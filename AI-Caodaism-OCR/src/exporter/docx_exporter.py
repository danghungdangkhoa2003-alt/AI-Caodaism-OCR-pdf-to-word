"""DOCX exporter enforcing Unicode Times New Roman 13pt output."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.pdf.pdf_processor import ContentBlock, PageContent


class DocumentExporter:
    """Create an editable Word document, using PDF block order as layout guidance."""

    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(13)

    def export(self, pages: list[PageContent], destination: Path, include_page_reference: bool = False) -> Path:
        document = Document()
        self._set_normal_style(document)
        for index, page in enumerate(pages):
            if index:
                document.add_section(WD_SECTION.NEW_PAGE)
            self._configure_page(document.sections[-1], page)
            self._write_page(document, page, include_page_reference)
        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
        return destination

    def _set_normal_style(self, document: Document) -> None:
        style = document.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)

    @staticmethod
    def _configure_page(section, page: PageContent) -> None:
        section.page_width = Pt(page.width)
        section.page_height = Pt(page.height)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    def _write_page(self, document: Document, page: PageContent, include_page_reference: bool) -> None:
        for block in page.blocks:
            if block.kind == "image" and block.image_path:
                self._add_image(document, block, page)
            elif block.kind == "text":
                self._add_text(document, block, page)
        if include_page_reference and page.rendered_page:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run("Bản quét gốc để đối chiếu").italic = True
            document.add_picture(str(page.rendered_page), width=Inches(6.2))

    def _add_text(self, document: Document, block: ContentBlock, page: PageContent) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
        width = block.bbox[2] - block.bbox[0]
        if width > page.width * 0.72:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif block.bbox[0] > page.width * 0.28:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for span in block.spans:
            run = paragraph.add_run(span.text)
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
            run.bold = span.bold
            run.italic = span.italic
        self._set_keep_lines(paragraph)

    @staticmethod
    def _add_image(document: Document, block: ContentBlock, page: PageContent) -> None:
        available_width = page.width - 2 * 0.65 * 72
        image_width = max(36, min(block.bbox[2] - block.bbox[0], available_width))
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            paragraph.add_run().add_picture(str(block.image_path), width=Pt(image_width))
        except UnrecognizedImageError:
            paragraph.add_run("[Không thể nhúng ảnh gốc; xem PDF để đối chiếu]").italic = True

    @staticmethod
    def _set_keep_lines(paragraph) -> None:
        properties = paragraph._p.get_or_add_pPr()
        keep_next = OxmlElement("w:keepNext")
        properties.append(keep_next)
